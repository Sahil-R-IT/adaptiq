"""
File content extraction for images and documents.
Supports: PDF, TXT, DOCX, and images via Gemini Vision.
"""
import os
import re
from pathlib import Path
from typing import Any, Optional

MAX_TEXT_CHARS = 12000

# --- PDF ---
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2 as pypdf
        PYPDF_AVAILABLE = True
    except ImportError:
        PYPDF_AVAILABLE = False

# --- DOCX ---
try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    try:
        from docx import Document as _DocxDocument
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False

ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "gif"}
ALLOWED_DOCUMENT_EXTENSIONS = {"pdf", "txt", "docx", "doc", "md"}


def _truncate(text: str, max_chars: int = MAX_TEXT_CHARS) -> str:
    if len(text) <= max_chars:
        return text.strip()
    truncated = text[:max_chars]
    last_period = truncated.rfind(".")
    if last_period > max_chars * 0.8:
        truncated = truncated[:last_period + 1]
    return truncated.strip() + "\n\n[Content truncated for quiz generation]"


def extract_pdf(file_path: str) -> str:
    if not PYPDF_AVAILABLE:
        raise RuntimeError("pypdf is not installed. Run: pip install pypdf")

    path = Path(file_path)
    if not path.exists():
        raise RuntimeError(f"File not found: {file_path}")

    try:
        reader = pypdf.PdfReader(str(path))
    except Exception as exc:
        raise RuntimeError(f"Could not open PDF: {exc}")

    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            pages.append(text)
        except Exception:
            continue

    combined = "\n\n".join(p.strip() for p in pages if p.strip())
    combined = re.sub(r"\s+", " ", combined).strip()
    combined = re.sub(r"[ \t]+\n", "\n", combined)

    if not combined or len(combined) < 20:
        raise RuntimeError(
            "No extractable text found in this PDF. "
            "It may be a scanned image PDF. Try uploading the PDF with selectable text."
        )

    return _truncate(combined)


def extract_txt(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise RuntimeError(f"File not found: {file_path}")

    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            text = path.read_text(encoding=encoding)
            text = re.sub(r"\r\n|\r", "\n", text)
            text = re.sub(r"\n{4,}", "\n\n\n", text)
            text = re.sub(r"[ \t]+", " ", text).strip()
            if text:
                return _truncate(text)
        except UnicodeDecodeError:
            continue

    raise RuntimeError("Could not read text file — unknown encoding.")


def extract_docx(file_path: str) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise RuntimeError(f"File not found: {file_path}")

    try:
        if hasattr(python_docx, "Document"):
            doc = python_docx.Document(str(path))
        else:
            from docx import Document
            doc = Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"Could not open DOCX: {exc}")

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    text = re.sub(r"\s+", " ", text).strip()

    if not text or len(text) < 20:
        raise RuntimeError("No extractable text found in this DOCX file.")

    return _truncate(text)


def extract_image_with_gemini(
    file_path: str,
    mime_type: str,
    gemini_api_key: str,
    gemini_model: str = "gemini-2.0-flash",
) -> str:
    """
    Use Gemini Vision to describe/extract content from an image.
    Returns a text description suitable for quiz generation.
    """
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        raise RuntimeError("Google GenAI SDK not installed. Run: pip install google-genai")

    if not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY is not set.")

    path = Path(file_path)
    if not path.exists():
        raise RuntimeError(f"Image file not found: {file_path}")

    image_bytes = path.read_bytes()

    client = genai.Client(api_key=gemini_api_key)

    prompt = (
        "You are analyzing an image to extract educational content for quiz generation. "
        "Describe all text, diagrams, charts, formulas, labels, data, and key concepts visible in this image. "
        "Be thorough and structured. If there is text in the image, transcribe it accurately. "
        "If there are diagrams or charts, describe what they show in detail. "
        "Output plain descriptive text only — no markdown, no bullet points, just clear prose."
    )

    try:
        response = client.models.generate_content(
            model=gemini_model,
            contents=[
                types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
                prompt,
            ],
        )
        description = (getattr(response, "text", "") or "").strip()
    except Exception as exc:
        raise RuntimeError(f"Gemini image analysis failed: {exc}")

    if not description or len(description) < 30:
        raise RuntimeError(
            "Gemini could not extract meaningful content from this image. "
            "Try an image with more visible text or diagrams."
        )

    return _truncate(description)


def extract_document(file_path: str, filename: str) -> dict[str, Any]:
    """
    Extract text from a document file.

    Returns:
        {
            "text": str,
            "title": str,
            "char_count": int,
            "file_name": str,
        }
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    title = filename

    if ext == "pdf":
        text = extract_pdf(file_path)
    elif ext in ("txt", "md"):
        text = extract_txt(file_path)
    elif ext in ("docx", "doc"):
        text = extract_docx(file_path)
    else:
        raise RuntimeError(f"Unsupported document type: .{ext}")

    return {
        "text": text,
        "title": title,
        "char_count": len(text),
        "file_name": filename,
    }


def get_mime_type(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return {
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "png": "image/png",
        "webp": "image/webp",
        "gif": "image/gif",
    }.get(ext, "image/jpeg")
